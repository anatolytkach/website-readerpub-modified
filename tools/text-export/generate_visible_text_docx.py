#!/usr/bin/env python3
from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import html


ROOT = Path(__file__).resolve().parents[2]
DIST = ROOT / "dist"
EXPORTS = ROOT / "exports"
OUT_DOCX = EXPORTS / "readerpub-visible-text-map.docx"
OUT_JSON = ROOT / "exports" / "text-map" / "readerpub-visible-text-map.json"

PAGE_ORDER = [
    ("home", "/", DIST / "index.html"),
    ("readers", "/readers/", DIST / "readers" / "index.html"),
    ("authors", "/authors/", DIST / "authors" / "index.html"),
    ("institutions", "/institutions/", DIST / "institutions" / "index.html"),
    ("platform", "/platform/", DIST / "platform" / "index.html"),
    ("weread", "/weread/", DIST / "weread" / "index.html"),
    ("wepub", "/wepub/", DIST / "wepub" / "index.html"),
    ("webuzz", "/webuzz/", DIST / "webuzz" / "index.html"),
    ("booktree", "/booktree/", DIST / "booktree" / "index.html"),
    ("security", "/security/", DIST / "security" / "index.html"),
    ("pricing", "/pricing/", DIST / "pricing" / "index.html"),
    ("contact", "/contact/", DIST / "contact" / "index.html"),
    ("kb", "/kb/", DIST / "kb" / "index.html"),
    ("start", "/start/", DIST / "start" / "index.html"),
    ("terms", "/terms/", DIST / "terms" / "index.html"),
    ("contacts-alias", "/contacts/", DIST / "contacts" / "index.html"),
    ("platform1-alias", "/platform1/", DIST / "platform1" / "index.html"),
    ("technology-alias", "/technology/", DIST / "technology" / "index.html"),
]

TEXT_TAGS = {
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "p",
    "li",
    "a",
    "button",
    "label",
    "legend",
    "summary",
    "figcaption",
    "option",
}

SKIP_TAGS = {"script", "style", "svg", "path", "img", "source", "picture", "meta", "link"}


@dataclass
class TextRecord:
    marker: str
    page_id: str
    route: str
    tag: str
    role: str
    text: str
    dom_path: str


def normalize_space(value: str) -> str:
    value = value.replace("\xa0", " ")
    value = re.sub(r"\s+", " ", value)
    return value.strip()


def element_text(el) -> str:
    return normalize_space(" ".join(el.itertext()))


def direct_text(el) -> str:
    parts = []
    if el.text:
        parts.append(el.text)
    for child in el:
        if child.tail:
            parts.append(child.tail)
    return normalize_space(" ".join(parts))


def should_skip(el) -> bool:
    tag = (el.tag or "").lower()
    if tag in SKIP_TAGS:
        return True
    if el.get("hidden") is not None:
        return True
    if el.get("aria-hidden") == "true" and tag not in {"div", "ul"}:
        return True
    if el.get("type") == "hidden":
        return True
    return False


def has_text_tag_child(el) -> bool:
    for child in el:
        child_tag = (child.tag or "").lower()
        if child_tag in TEXT_TAGS:
            return True
    return False


def dom_path(el) -> str:
    parts = []
    cur = el
    while cur is not None and getattr(cur, "tag", None) is not None:
        tag = str(cur.tag).lower()
        if tag in {"html", "body"}:
            parts.append(tag)
            break
        parent = cur.getparent()
        index = 1
        if parent is not None:
            siblings = [sib for sib in parent if getattr(sib, "tag", None) == cur.tag]
            if len(siblings) > 1:
                index = siblings.index(cur) + 1
                parts.append(f"{tag}[{index}]")
            else:
                parts.append(tag)
        else:
            parts.append(tag)
        cur = parent
    return " > ".join(reversed(parts))


def role_for(el) -> str:
    classes = el.get("class", "")
    tag = (el.tag or "").lower()
    if tag.startswith("h"):
        return f"heading {tag[-1]}"
    if "button" in classes or tag == "button":
        return "button"
    if tag == "a":
        return "link"
    if tag == "li":
        return "list item"
    if tag == "p":
        return "paragraph"
    return tag


def make_marker(page_id: str, index: int, text: str) -> str:
    digest = hashlib.sha1(text.encode("utf-8")).hexdigest()[:8]
    return f"TXT:{page_id}:{index:04d}:{digest}"


def extract_page(page_id: str, route: str, path: Path) -> list[TextRecord]:
    doc = html.fromstring(path.read_text(encoding="utf-8"))
    body = doc.find("body")
    if body is None:
        return []

    records: list[TextRecord] = []
    seen_marker_texts: set[tuple[str, str, str]] = set()

    for el in body.iter():
        tag = (el.tag or "").lower()
        if should_skip(el):
            continue

        text = ""
        if tag in TEXT_TAGS:
            if tag == "li" and any((child.tag or "").lower() == "ul" for child in el):
                text = direct_text(el)
            elif tag in {"a", "button", "label", "option"}:
                text = element_text(el)
            elif has_text_tag_child(el):
                continue
            else:
                text = element_text(el)

        if text:
            key = (tag, dom_path(el), text)
            if key not in seen_marker_texts:
                marker = make_marker(page_id, len(records) + 1, text)
                records.append(
                    TextRecord(
                        marker=marker,
                        page_id=page_id,
                        route=route,
                        tag=tag,
                        role=role_for(el),
                        text=text,
                        dom_path=dom_path(el),
                    )
                )
                seen_marker_texts.add(key)

        if tag in {"input", "textarea"}:
            placeholder = normalize_space(el.get("placeholder", ""))
            value = normalize_space(el.get("value", ""))
            label_text = placeholder or value
            if label_text:
                marker = make_marker(page_id, len(records) + 1, label_text)
                records.append(
                    TextRecord(
                        marker=marker,
                        page_id=page_id,
                        route=route,
                        tag=tag,
                        role="form placeholder/value",
                        text=label_text,
                        dom_path=dom_path(el),
                    )
                )

    return records


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_run(paragraph, text: str, *, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)
    return run


def build_docx(pages: dict[str, list[TextRecord]]) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Title", 20, "000000"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)

    title = document.add_paragraph()
    title.style = styles["Title"]
    add_run(title, "ReaderPub Modified: visible text map", bold=True)

    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    add_run(intro, "Instructions: ", bold=True)
    intro.add_run(
        "edit only the text after each marker. Keep markers like [[TXT:page:0001:hash]] unchanged. "
        "If you need to delete a text fragment, leave the marker and write DELETE after it."
    )

    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}. ")
    meta.add_run("Source: built static HTML in /dist.")

    for page_num, (page_id, route, _) in enumerate(PAGE_ORDER, start=1):
        records = pages.get(page_id, [])
        if not records:
            continue
        if page_num > 1:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        h1 = document.add_heading(f"{route}  [{page_id}]", level=1)
        h1.paragraph_format.keep_with_next = True

        summary = document.add_paragraph()
        add_run(summary, f"{len(records)} text fragments", italic=True, color="555555")

        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.autofit = False
        widths = [2500, 1300, 5560]
        headers = ["Marker", "Role", "Editable text"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            set_cell_width(cell, widths[i])
            set_cell_shading(cell, "E8EEF5")
            paragraph = cell.paragraphs[0]
            add_run(paragraph, header, bold=True)

        for record in records:
            row = table.add_row()
            cells = row.cells
            for i, width in enumerate(widths):
                set_cell_width(cells[i], width)

            marker_para = cells[0].paragraphs[0]
            add_run(marker_para, f"[[{record.marker}]]", bold=True, color="1F4D78", size=8.5)
            marker_para.add_run(f"\n{record.tag}")

            role_para = cells[1].paragraphs[0]
            role_para.add_run(record.role)

            text_para = cells[2].paragraphs[0]
            text_para.add_run(record.text)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    OUT_JSON.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT_DOCX)


def main() -> None:
    pages: dict[str, list[TextRecord]] = {}
    missing = []
    for page_id, route, path in PAGE_ORDER:
        if not path.exists():
            missing.append(str(path))
            continue
        pages[page_id] = extract_page(page_id, route, path)

    data = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "source": "dist static HTML",
        "docx": str(OUT_DOCX.relative_to(ROOT)),
        "pages": {page_id: [asdict(record) for record in records] for page_id, records in pages.items()},
        "missing": missing,
    }
    OUT_JSON.parent.mkdir(parents=True, exist_ok=True)
    OUT_JSON.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    build_docx(pages)
    print(OUT_DOCX)
    print(OUT_JSON)
    print(sum(len(records) for records in pages.values()), "text fragments")


if __name__ == "__main__":
    main()
